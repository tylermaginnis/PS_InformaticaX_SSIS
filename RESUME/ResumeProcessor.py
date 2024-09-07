from docx import Document
import re

def process_resume(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    document = Document()
    in_key_skills_section = False
    in_experience_section = False
    in_duties_section = False
    key_skills_paragraph = None
    for line in lines:
        line = line.replace('â€“', '–')  # Replace incorrect character with emdash
        if line.startswith('<skills>'):
            in_key_skills_section = True
            key_skills_paragraph = document.add_paragraph()
        elif line.startswith('</skills>'):
            in_key_skills_section = False
        elif line.startswith('<experience>'):
            in_experience_section = True
        elif line.startswith('</experience>'):
            in_experience_section = False
        elif line.startswith('<duties>') and in_experience_section:
            in_duties_section = True
        elif line.startswith('</duties>') and in_experience_section:
            in_duties_section = False
        elif line.startswith('<job>') and in_experience_section:
            job_paragraph = document.add_paragraph()
            job_paragraph.style = 'List Bullet'
            run = job_paragraph.add_run(line[5:-6])
            run.bold = True
        elif line.startswith('</job>') and in_experience_section:
            continue  # Skip the </job> tag
        elif line.startswith('<skill>') and in_key_skills_section:
            skill_text = line[7:].strip().replace('</skill>', '')
            parts = re.split(r'(\*\*.*?\*\*)', skill_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = key_skills_paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = key_skills_paragraph.add_run(part)
            key_skills_paragraph.add_run().add_break()  # Shift+return newline
        elif line.startswith('<header>'):
            document.add_paragraph()  # Linebreak before header
            run = document.add_paragraph().add_run(line[8:].strip())
            run.bold = True
            document.add_paragraph()  # Linebreak after header
        elif line.startswith('</header>'):
            continue  # Skip the </header> tag
        elif line.startswith('<objective>') or line.startswith('</objective>'):
            continue  # Skip the <objective> and </objective> tags
        else:
            if in_key_skills_section:
                parts = re.split(r'(\*\*.*?\*\*)', line.strip().replace('</skill>', ''))
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = key_skills_paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = key_skills_paragraph.add_run(part)
                key_skills_paragraph.add_run(' ')
                key_skills_paragraph.paragraph_format.space_after = 0  # Eliminate line spacing in KEY SKILLS section
            elif in_duties_section:
                if line.strip().startswith('-'):
                    duty_paragraph = document.add_paragraph()
                    duty_paragraph.style = 'List Bullet'
                    run = duty_paragraph.add_run(line.strip()[1:].strip())
                else:
                    run = duty_paragraph.add_run(' ' + line.strip())
            else:
                if line.strip():  # Only add non-empty lines
                    paragraph = document.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*|_.*?_)', line.strip())
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = paragraph.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('_') and part.endswith('_'):
                            run = paragraph.add_run(part[1:-1])
                            run.italic = True
                        else:
                            run = paragraph.add_run(part)
                    paragraph.paragraph_format.space_after = 0  # Equivalent to shift+enter in Word

    document.save('Processed_Resume.docx')

if __name__ == "__main__":
    process_resume('Resume.txt')
