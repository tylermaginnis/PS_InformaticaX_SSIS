from docx import Document
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.shared import RGBColor

def process_coverletter(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    document = Document()
    in_coverletter_section = False
    for line in lines:
        line = line.replace('â€“', '–')  # Replace incorrect character with emdash
        if line.startswith('<coverletter>'):
            in_coverletter_section = True
        elif line.startswith('</coverletter>'):
            in_coverletter_section = False
        elif line.startswith('<header>'):
            document.add_paragraph()  # Linebreak before header
            run = document.add_paragraph().add_run(line[8:].strip())
            run.bold = True
            document.add_paragraph()  # Linebreak after header
        elif line.startswith('</header>'):
            continue  # Skip the </header> tag
        else:
            if in_coverletter_section:
                if line.strip():  # Only add non-empty lines
                    paragraph = document.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*|_.*?_|!?\[.*?\]\(.*?\))', line.strip())
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = paragraph.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('_') and part.endswith('_'):
                            run = paragraph.add_run(part[1:-1])
                            run.italic = True
                        elif part.startswith('[') and '](' in part and part.endswith(')'):
                            link_text = re.search(r'\[(.*?)\]', part).group(1)
                            link_url = re.search(r'\((.*?)\)', part).group(1)
                            run = paragraph.add_run(f"{link_text} ({link_url})")
                            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
                            run.font.underline = True
                            # Create hyperlink
                            hyperlink = OxmlElement('w:hyperlink')
                            hyperlink.set(qn('r:id'), document.part.relate_to(link_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
                            hyperlink.append(run._r)
                            paragraph._element.append(hyperlink)  # Append hyperlink to the paragraph element
                        else:
                            run = paragraph.add_run(part)
                    paragraph.paragraph_format.space_after = 0  # Equivalent to shift+enter in Word

    document.save('Processed_Coverletter.docx')

if __name__ == "__main__":
    process_coverletter('Coverletter.txt')
